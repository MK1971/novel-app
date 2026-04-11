#!/usr/bin/env python3
"""Generate docs/Novel_App_SSH_Operations_and_Security_Review.docx — run from repo root."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def main():
    out = Path(__file__).resolve().parent / "Novel_App_SSH_Operations_and_Security_Review.docx"
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Novel App — SSH operations & security reference")
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph(
        "Generated for operator use. Fill in placeholders from Cloudways (Access Details) per environment. "
        "Do not commit real passwords or API keys into version control."
    )

    add_heading(doc, "1. Environment map (from project handbook)", 1)
    doc.add_paragraph(
        "Branch → hostname (typical):\n"
        "• Development → dev.whatsmybookname.com\n"
        "• staging → staging.whatsmybookname.com\n"
        "• Production → www.whatsmybookname.com\n"
        "Rule: dev and staging should stay private (e.g. Cloudflare Access); production is public."
    )

    add_heading(doc, "2. SSH: how to connect (all three environments)", 1)
    doc.add_paragraph(
        "Cloudways gives each application its own SSH user (master_…) and each server has an IP. "
        "Copy the exact user, host, port, and application path from: "
        "Cloudways → your Server → Application → Access Details (or Master Credentials)."
    )
    add_heading(doc, "2.1 SSH key (local machine)", 2)
    doc.add_paragraph(
        "Use the private key you added in Cloudways → SSH Keys (e.g. ~/.ssh/cloudways_wmbn_prod). "
        "Never share the private key or commit it to git."
    )
    add_heading(doc, "2.2 Connection template", 2)
    p = doc.add_paragraph()
    p.add_run("ssh -i ~/.ssh/<YOUR_KEY> -p 22 <MASTER_USER>@<SERVER_IP>\n").font.name = "Courier New"
    doc.add_paragraph("Examples (replace with your Access Details):")
    for label, note in [
        ("Dev", "SSH user + IP from the dev application’s Access Details."),
        ("Staging", "SSH user + IP from the staging application’s Access Details (may differ from dev)."),
        ("Production", "SSH user + IP from the production application’s Access Details."),
    ]:
        doc.add_paragraph(f"{label}: {note}", style="List Bullet")

    add_heading(doc, "2.3 Application root (Laravel) on the server", 2)
    doc.add_paragraph(
        "After SSH, the Laravel project root is usually:\n"
        "~/applications/<APP_FOLDER_NAME>/public_html\n"
        "Web server document root is typically public_html/public.\n"
        "The folder name (e.g. cuwvnwmwrf) is unique per app — confirm in Access Details."
    )

    add_heading(doc, "3. Commands we used and when to use them again", 1)

    rows = [
        (
            "cd ~/applications/<APP_FOLDER>/public_html",
            "Change to the Laravel root before running artisan.",
            "Every SSH session when running PHP/Artisan on that app.",
        ),
        (
            "php artisan config:clear",
            "Clears cached configuration so Laravel will re-read .env on next request (when not using only file cache).",
            "After changing .env keys such as ADMIN_EMAIL, APP_URL, mail, OAuth, PayPal.",
        ),
        (
            "php artisan config:cache",
            "Rebuilds bootstrap/cache/config.php from .env and config files. Required in production after .env changes.",
            "After any production .env edit; then verify with config:show (below).",
        ),
        (
            "php artisan config:show app.admin_email",
            "Shows the resolved admin email used by the admin Gate (must match your operator account email if using ADMIN_EMAIL).",
            "Troubleshooting admin access; after setting ADMIN_EMAIL in .env.",
        ),
        (
            "php artisan optimize:clear",
            "Clears config, route, view, and other bootstrap caches.",
            "After deploys or when routes/views behave stale.",
        ),
        (
            "php artisan route:cache",
            "Caches routes for faster requests.",
            "Production after route changes; run after optimize:clear if you use route caching.",
        ),
        (
            "php artisan view:cache",
            "Compiles Blade to cache.",
            "Production after view changes.",
        ),
        (
            "php artisan migrate --force",
            "Runs database migrations (non-interactive).",
            "Deploys that ship migrations. Use --force only in staging/production as appropriate.",
        ),
        (
            "php artisan db:seed --class=Database\\\\Seeders\\\\AdminOnlySeeder",
            "Creates/updates the seeded admin user (password default in seeder — change immediately in prod).",
            "Rare: initial admin bootstrap; prefer ADMIN_EMAIL + real account when possible.",
        ),
        (
            "grep ^ADMIN_EMAIL .env",
            "Quick check that ADMIN_EMAIL exists in the active .env on the server.",
            "Debugging admin email / config mismatches.",
        ),
        (
            "nano .env",
            "Edit environment variables on the server.",
            "Adding ADMIN_EMAIL, fixing APP_URL, mail, keys.",
        ),
    ]
    for cmd, purpose, again in rows:
        p = doc.add_paragraph()
        p.add_run("Command:\n").bold = True
        p.add_run(cmd + "\n").font.name = "Courier New"
        p.add_run("Purpose: ").bold = True
        p.add_run(purpose + "\n")
        p.add_run("When again: ").bold = True
        p.add_run(again)

    add_heading(doc, "4. Deploying a single file from your laptop (SCP)", 1)
    doc.add_paragraph(
        "Used when you patch one PHP file without a full Git deploy. "
        "Replace local path, key, user, host, and remote path."
    )
    p = doc.add_paragraph()
    p.add_run(
        "scp -i ~/.ssh/<YOUR_KEY> -P 22 \\\n"
        "  /path/to/local/app/Http/Controllers/SomeController.php \\\n"
        "  '<MASTER_USER>@<SERVER_IP>:~/applications/<APP_FOLDER>/public_html/app/Http/Controllers/'\n"
    ).font.name = "Courier New"
    doc.add_paragraph(
        "After copying PHP or config files, run on the server: php artisan optimize:clear && php artisan config:cache && php artisan route:cache && php artisan view:cache (adjust as needed)."
    )

    add_heading(doc, "5. Security review (application — static assessment)", 1)
    doc.add_paragraph(
        "This is a code-and-configuration review, not an authorized live penetration test. "
        "Findings below are recommendations; prioritize by risk for your threat model."
    )

    findings = [
        (
            "Dependency advisories (Composer)",
            "Run composer audit regularly (composer audit). Keep league/commonmark and phpseclib (and other "
            "transitive deps) on patched releases; re-run tests after composer update.",
        ),
        (
            "Secrets and .env",
            "Never commit .env or scripts containing live credentials (e.g. prod_secrets.local.sh should stay gitignored). "
            "Rotate any secret that may have been copied into tickets, chat, or screenshots.",
        ),
        (
            "Admin access",
            "ADMIN_EMAIL and is_admin: prefer config:cache after .env changes; use config(app.admin_email) pattern (implemented in codebase). "
            "Restrict who can register the ADMIN_EMAIL address.",
        ),
        (
            "Debug mode",
            "Ensure APP_DEBUG=false and APP_ENV=production on production.",
        ),
        (
            "Sessions / HTTPS",
            "SESSION_SECURE_COOKIE=true behind HTTPS; Cloudflare SSL mode Full (strict) with valid origin cert recommended.",
        ),
        (
            "Rate limiting",
            "Login uses throttling (LoginRequest / auth routes). Payment checkout uses RateLimiter. "
            "Consider additional throttling on sensitive POST endpoints if abuse appears.",
        ),
        (
            "Origin exposure",
            "If traffic should only come from Cloudflare, restrict the origin firewall to Cloudflare IP ranges "
            "and/or use Authenticated Origin Pulls (Cloudflare) so direct-to-IP bypass is harder.",
        ),
        (
            "Cloudflare Access",
            "Dev/staging behind Access reduces scanning of non-prod apps; keep production DNS and Access policies separate.",
        ),
        (
            "PayPal / webhooks",
            "Verify webhook signatures and idempotency in PaymentController flows; keep PAYPAL_* and webhook tokens secret.",
        ),
        (
            "OAuth",
            "Google redirect URIs must match exactly; register apex and www if both are used.",
        ),
    ]
    for title, body in findings:
        p = doc.add_paragraph()
        p.add_run(title + ": ").bold = True
        p.add_run(body)

    add_heading(doc, "6. Penetration testing — scope and ethics", 1)
    doc.add_paragraph(
        "Unauthorized testing against systems you do not own or lack written permission to test is illegal and unethical. "
        "For production, use one or more of:"
    )
    doc.add_paragraph(
        "• Written scope with your hosting/DNS providers if needed.\n"
        "• A qualified third-party pentest with rules of engagement.\n"
        "• Bug bounty program if you run one (clear scope and safe harbor).\n"
        "• Internal testing: staging clone with anonymized data, not live PII.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Automated and manual checks you can run safely on your own staging instance:\n"
        "• OWASP ZAP or Burp Suite (authenticated crawl of staging).\n"
        "• composer audit / npm audit for dependency CVEs.\n"
        "• Laravel: php artisan route:list review for unintended public routes.\n"
        "• Verify CSRF on state-changing forms; verify authorization on IDOR-prone routes (profile, edits, payments).\n"
        "• SSL Labs / securityheaders.com for TLS and headers (staging/prod URLs)."
    )

    add_heading(doc, "7. Revision log", 1)
    doc.add_paragraph(
        "Document generated to consolidate SSH usage from deployment work: "
        "config caching, ADMIN_EMAIL troubleshooting, single-file SCP deploys, and Cloudflare Access for non-prod. "
        "Update hostnames, paths, and commands when Cloudways or DNS changes."
    )

    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
